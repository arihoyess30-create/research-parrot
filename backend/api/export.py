import json, io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from db import get_db
from models import Project, User
from api.auth import get_current_user

router = APIRouter()

def require_paid(proj):
    if not proj.is_paid:
        raise HTTPException(402, "Payment required to export")

def get_ordered_sections(proj):
    outline  = json.loads(proj.outline)  if proj.outline  else {"chapters": []}
    sections = json.loads(proj.sections) if proj.sections else {}
    chapters = outline.get("chapters", [])
    result   = [(f"Chapter {ch['number']}: {ch['title']}",
                 sections.get(str(ch["number"]), "[Content not yet generated]"))
                for ch in chapters]
    return result, outline.get("title", proj.title)

def build_pdf(proj) -> io.BytesIO:
    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(buf, pagesize=A4,
                              leftMargin=3*cm, rightMargin=2.5*cm,
                              topMargin=2.5*cm, bottomMargin=2.5*cm)
    ss   = getSampleStyleSheet()
    title_style = ParagraphStyle("rp_title", parent=ss["Title"], fontSize=18,
        textColor=colors.HexColor("#0F6E56"), alignment=TA_CENTER, spaceAfter=6)
    sub_style   = ParagraphStyle("rp_sub", parent=ss["Normal"], fontSize=11,
        textColor=colors.HexColor("#5F5E5A"), alignment=TA_CENTER, spaceAfter=4)
    chap_style  = ParagraphStyle("rp_chap", parent=ss["Heading1"], fontSize=14,
        textColor=colors.HexColor("#0F6E56"), spaceBefore=24, spaceAfter=10)
    body_style  = ParagraphStyle("rp_body", parent=ss["Normal"], fontSize=11,
        leading=18, alignment=TA_JUSTIFY, spaceAfter=10)

    sections, paper_title = get_ordered_sections(proj)
    elements = [
        Spacer(1, 3*cm),
        Paragraph(paper_title, title_style),
        Spacer(1, .5*cm),
        Paragraph(f"Field: {proj.field}  ·  Level: {proj.level.title()}  ·  {proj.citation_style}", sub_style),
        Paragraph("Research Parrot — AI Academic Research Assistant", sub_style),
        Spacer(1, 1*cm),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1D9E75")),
        PageBreak(),
    ]
    for chap_title, content in sections:
        elements.append(Paragraph(chap_title, chap_style))
        for para in content.split("\n\n"):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
        elements.append(PageBreak())
    doc.build(elements)
    buf.seek(0)
    return buf

def build_docx(proj) -> io.BytesIO:
    doc = DocxDocument()
    sections_data, paper_title = get_ordered_sections(proj)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(0.98)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(paper_title)
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    doc.add_paragraph(f"Field: {proj.field}  |  Level: {proj.level.title()}  |  {proj.citation_style}") \
       .alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Research Parrot — AI Academic Research Assistant") \
       .alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    for chap_title, content in sections_data:
        h = doc.add_heading(chap_title, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
        for para in content.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)
        doc.add_page_break()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

@router.get("/pdf/{project_id}")
def export_pdf(project_id: str, db: Session = Depends(get_db),
               user: User = Depends(get_current_user)):
    proj = db.query(Project).filter_by(id=project_id, user_id=user.id).first()
    if not proj: raise HTTPException(404, "Project not found")
    require_paid(proj)
    buf  = build_pdf(proj)
    name = proj.title[:60].replace(" ", "_") + ".pdf"
    return StreamingResponse(buf, media_type="application/pdf",
                             headers={"Content-Disposition": f'attachment; filename="{name}"'})

@router.get("/docx/{project_id}")
def export_docx(project_id: str, db: Session = Depends(get_db),
                user: User = Depends(get_current_user)):
    proj = db.query(Project).filter_by(id=project_id, user_id=user.id).first()
    if not proj: raise HTTPException(404, "Project not found")
    require_paid(proj)
    buf  = build_docx(proj)
    name = proj.title[:60].replace(" ", "_") + ".docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return StreamingResponse(buf, media_type=mime,
                             headers={"Content-Disposition": f'attachment; filename="{name}"'})
